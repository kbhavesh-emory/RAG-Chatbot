"""Common ingestion utilities shared by app + offline embedding script.

This module intentionally has no Streamlit dependency so it can be imported from
CLI scripts without triggering UI code.
"""

from __future__ import annotations

import json
import os
import shutil
import subprocess
import tempfile
import uuid
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Iterable, List, Optional


# Optional deps (installed via requirements.txt in this repo)
try:
    import pandas as pd
except Exception:  # pragma: no cover
    pd = None

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover
    DocxDocument = None


SUPPORTED_EXTENSIONS = {"pdf", "txt", "docx", "doc", "xlsx", "xls"}


@dataclass
class SimpleDocument:
    """Minimal document object compatible with LangChain vector stores."""

    page_content: str
    metadata: dict
    id: str = field(default_factory=lambda: str(uuid.uuid4()))


class SimpleTextSplitter:
    """Simple text splitter."""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_text(self, text: str) -> List[str]:
        chunks: List[str] = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size

            if end < len(text):
                for delimiter in [". ", "\n\n", "\n", " "]:
                    last_pos = text.rfind(delimiter, start, end)
                    if last_pos > start:
                        end = last_pos + len(delimiter)
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - self.chunk_overlap

        return chunks

    def split_documents(self, documents: Iterable[SimpleDocument]) -> List[SimpleDocument]:
        split_docs: List[SimpleDocument] = []

        for doc in documents:
            text_chunks = self.split_text(doc.page_content)

            for i, chunk in enumerate(text_chunks):
                metadata = dict(doc.metadata)
                metadata["chunk"] = i
                split_docs.append(
                    SimpleDocument(
                        page_content=chunk,
                        metadata=metadata,
                        id=f"{doc.id or 'doc'}_chunk_{i}",
                    )
                )

        return split_docs


def _parse_optional_limit(value: Optional[str], default_value: int | None) -> int | None:
    if value is None:
        return default_value

    raw = value.strip()
    if raw == "":
        return default_value

    lowered = raw.lower()
    if lowered in {
        "unlimited",
        "unlomited",  # common typo
        "none",
        "null",
        "inf",
        "infinite",
        "no-limit",
        "nolimit",
    }:
        return None

    parsed = int(raw)
    if parsed <= 0:
        return None

    return parsed


def load_docx(path: str, source_name: str) -> List[SimpleDocument]:
    if DocxDocument is None:
        raise RuntimeError("Missing dependency: python-docx. Install it with: pip install python-docx")

    doc = DocxDocument(path)
    parts: List[str] = []

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_values = [((cell.text or "").strip()) for cell in row.cells]
            row_values = [v for v in row_values if v]
            if row_values:
                parts.append("\t".join(row_values))

    full_text = "\n".join(parts).strip() or "(No extractable text found in DOCX.)"

    return [
        SimpleDocument(
            page_content=full_text,
            metadata={"source": source_name, "file_type": "docx"},
            id=str(uuid.uuid4()),
        )
    ]


def load_excel(
    path: str,
    source_name: str,
    *,
    excel_max_rows: int | None = 2000,
    excel_max_cols: int | None = 50,
) -> List[SimpleDocument]:
    if pd is None:
        raise RuntimeError("Missing dependency: pandas. Install it with: pip install pandas openpyxl xlrd")

    sheets = pd.read_excel(path, sheet_name=None)
    out: List[SimpleDocument] = []

    for sheet_name, df in sheets.items():
        if df is None:
            continue

        truncated = False
        if excel_max_rows is not None and df.shape[0] > excel_max_rows:
            df = df.head(excel_max_rows)
            truncated = True
        if excel_max_cols is not None and df.shape[1] > excel_max_cols:
            df = df.iloc[:, :excel_max_cols]
            truncated = True

        csv_text = df.to_csv(index=False)
        header = f"Sheet: {sheet_name}\n"
        if truncated:
            header += (
                f"(Truncated to {excel_max_rows if excel_max_rows is not None else 'unlimited'} rows "
                f"and {excel_max_cols if excel_max_cols is not None else 'unlimited'} cols)\n"
            )

        out.append(
            SimpleDocument(
                page_content=header + csv_text,
                metadata={"source": source_name, "file_type": "excel", "sheet": sheet_name},
                id=str(uuid.uuid4()),
            )
        )

    if not out:
        out.append(
            SimpleDocument(
                page_content="(No readable sheets found in Excel file.)",
                metadata={"source": source_name, "file_type": "excel"},
                id=str(uuid.uuid4()),
            )
        )

    return out


def load_doc_legacy(path: str, source_name: str) -> List[SimpleDocument]:
    # Best-effort support for .doc (legacy Word binary format).
    # 1) If LibreOffice exists, convert to .docx and parse.
    # 2) Else if antiword exists, extract plain text.

    if shutil.which("soffice"):
        with tempfile.TemporaryDirectory(prefix="rag_doc_convert_") as tmp:
            proc = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--nologo",
                    "--nolockcheck",
                    "--nodefault",
                    "--nofirststartwizard",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    tmp,
                    path,
                ],
                capture_output=True,
                text=True,
            )
            if proc.returncode != 0:
                raise RuntimeError(
                    "Failed to convert .doc via LibreOffice. "
                    f"stderr: {proc.stderr.strip() or proc.stdout.strip()}"
                )

            base = os.path.splitext(os.path.basename(path))[0]
            converted = os.path.join(tmp, base + ".docx")
            if not os.path.exists(converted):
                candidates = [
                    os.path.join(tmp, f)
                    for f in os.listdir(tmp)
                    if f.lower().endswith(".docx")
                ]
                if not candidates:
                    raise RuntimeError("LibreOffice conversion produced no .docx output")
                converted = candidates[0]

            return load_docx(converted, source_name)

    if shutil.which("antiword"):
        proc = subprocess.run(["antiword", path], capture_output=True, text=True)
        if proc.returncode != 0:
            raise RuntimeError(
                "Failed to extract .doc via antiword. "
                f"stderr: {proc.stderr.strip() or proc.stdout.strip()}"
            )
        text_out = (proc.stdout or "").strip() or "(No extractable text found in DOC.)"
        return [
            SimpleDocument(
                page_content=text_out,
                metadata={"source": source_name, "file_type": "doc"},
                id=str(uuid.uuid4()),
            )
        ]

    raise RuntimeError(
        "Legacy .doc requires LibreOffice (soffice) or antiword installed on the server. "
        "Prefer uploading .docx when possible."
    )


def file_extension(path: str) -> str:
    return os.path.splitext(path)[1].lower().lstrip(".")


def load_file_to_documents(
    path: str,
    *,
    excel_max_rows: int | None = 2000,
    excel_max_cols: int | None = 50,
) -> List[SimpleDocument]:
    """Load a single file into one-or-more SimpleDocuments."""

    ext = file_extension(path)
    source_name = os.path.basename(path)

    if ext not in SUPPORTED_EXTENSIONS:
        raise RuntimeError(f"Unsupported file type: .{ext}")

    if ext == "docx":
        return load_docx(path, source_name)

    if ext == "doc":
        return load_doc_legacy(path, source_name)

    if ext in ("xlsx", "xls"):
        return load_excel(path, source_name, excel_max_rows=excel_max_rows, excel_max_cols=excel_max_cols)

    # PDF/TXT: use LangChain community loaders
    from langchain_community.document_loaders import PyPDFLoader, TextLoader

    if ext == "pdf":
        loader = PyPDFLoader(path)
    else:
        loader = TextLoader(path)

    loaded_docs = loader.load()
    out: List[SimpleDocument] = []

    for doc in loaded_docs:
        out.append(
            SimpleDocument(
                page_content=doc.page_content,
                metadata={**(doc.metadata or {}), "source": source_name, "file_type": ext},
                id=str(uuid.uuid4()),
            )
        )

    return out


def list_upload_files(uploads_dir: str) -> List[str]:
    files: List[str] = []

    if not os.path.isdir(uploads_dir):
        return files

    for root, _, filenames in os.walk(uploads_dir):
        for name in filenames:
            ext = file_extension(name)
            if ext in SUPPORTED_EXTENSIONS:
                files.append(os.path.join(root, name))

    files.sort()
    return files


def write_manifest(
    index_dir: str,
    *,
    embedding_model: str,
    chunk_size: int,
    chunk_overlap: int,
    excel_max_rows: int | None,
    excel_max_cols: int | None,
    file_paths: List[str],
) -> None:
    os.makedirs(index_dir, exist_ok=True)

    manifest = {
        "created_at": datetime.now(timezone.utc).isoformat(),
        "embedding_model": embedding_model,
        "chunk_size": chunk_size,
        "chunk_overlap": chunk_overlap,
        "excel_max_rows": excel_max_rows,
        "excel_max_cols": excel_max_cols,
        "file_count": len(file_paths),
        "files": [os.path.basename(p) for p in file_paths],
    }

    with open(os.path.join(index_dir, "manifest.json"), "w", encoding="utf-8") as f:
        json.dump(manifest, f, indent=2)
